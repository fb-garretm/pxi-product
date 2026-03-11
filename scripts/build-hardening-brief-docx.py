"""
Build hardening-sprint-brief.docx with the same header style as PRDs.
Requires: python-docx, reference PRD docx, scripts/logo.jpg
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
REPO_ROOT = os.path.join(SCRIPT_DIR, "..")
REF_PATH = os.path.join(REPO_ROOT, "prds", "custom-domains-crawl.docx")
OUT_PATH = os.path.join(REPO_ROOT, "strategy", "hardening-sprint-brief.docx")
LOGO_PATH = os.path.join(SCRIPT_DIR, "logo.jpg")

if not os.path.isfile(REF_PATH):
    REF_PATH = "/Users/garretmann/Downloads/Swoop x Club App PRD.docx"

BLUE = RGBColor(0x3D, 0x85, 0xC6)
BEBAS = "Bebas Neue"


def add_title_block(doc):
    if os.path.isfile(LOGO_PATH):
        p = doc.add_paragraph()
        run = p.add_run()
        run.add_picture(LOGO_PATH, width=Inches(2.58), height=Inches(0.35))
    p = doc.add_paragraph()
    run = p.add_run("Hardening Sprint Brief\nEvent Directory, CMS & OSS")
    run.font.name = BEBAS
    run.font.size = Pt(48)
    run.font.bold = True
    run.font.color.rgb = BLUE
    p.style = doc.styles["Title"]
    p = doc.add_paragraph()
    run = p.add_run("Sprint Brief")
    run.font.name = BEBAS
    run.font.size = Pt(30)
    p = doc.add_paragraph()
    run = p.add_run("March 2026")
    run.font.size = Pt(10)
    p = doc.add_paragraph()
    run = p.add_run("Confidential")
    run.font.size = Pt(10)
    run.font.bold = True
    for _ in range(3):
        doc.add_paragraph()


def add_h1(doc, text):
    p = doc.add_paragraph(text, style="Heading 1")
    for run in p.runs:
        run.font.name = BEBAS
        run.font.color.rgb = BLUE


def add_h2(doc, text):
    p = doc.add_paragraph(text, style="Heading 2")
    for run in p.runs:
        run.font.name = BEBAS
        run.font.size = Pt(15)


def add_body(doc, text):
    doc.add_paragraph(text)


def add_bullet(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Inches(0.5)
    pf.first_line_indent = Inches(-0.25)
    p.add_run("\u2022  " + text)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.bold = True
        run.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
    tbl = table._element
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl.makeelement(qn("w:tblPr"), {})
    borders = tblPr.makeelement(qn("w:tblBorders"), {})
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.makeelement(qn(f"w:{edge}"), {
            qn("w:val"): "single",
            qn("w:sz"): "4",
            qn("w:space"): "0",
            qn("w:color"): "999999",
        })
        borders.append(el)
    tblPr.append(borders)
    return table


def main():
    doc = Document(REF_PATH)
    for table in list(doc.tables):
        table._element.getparent().remove(table._element)
    for p in list(doc.paragraphs):
        p._element.getparent().remove(p._element)

    add_title_block(doc)

    add_h1(doc, "Sprint overview")
    add_body(doc, "Hardening sprint for the recently developed public event directory, CMS, and Operator Standard Site (OSS). Focus: stability, polish, and final touches before release.")
    add_body(doc, "Related PRDs: Event Directory, CMS V0, OSS (see repo prds/).")
    doc.add_paragraph()

    add_h1(doc, "Sprint goals")
    add_bullet(doc, "Ship all P0 (must-have) items so Event Directory, CMS, and OSS are release-ready.")
    add_bullet(doc, "Address P1 (nice-to-have) polish and UX where capacity allows.")
    add_bullet(doc, "No critical regressions; shared surfaces (event cards, schedule, brackets) stable across Event Directory and OSS.")
    doc.add_paragraph()

    add_h1(doc, "Definition of Done")
    add_bullet(doc, "All P0 tickets merged and QA-verified (or explicitly deferred with agreement).")
    add_bullet(doc, "No critical or high-severity regressions on Event Directory, CMS, or OSS.")
    add_bullet(doc, "Staging deployed with hardening changes; release cadence (code freeze \u2192 QA \u2192 staging \u2192 production) respected per OSS/Event Directory timelines.")
    doc.add_paragraph()

    add_h1(doc, "P0 \u2014 Must-haves (Highest priority)")
    add_body(doc, "8 tickets. Complete before sprint close.")
    add_table(doc, ["Key", "Summary", "Assignee", "Status"], [
        ["SWA-170", "API Change from fb_compete.sports \u2192 public.sports", "\u2014", "Backlog"],
        ["SWA-169", "Create CMS Settings Page", "\u2014", "Backlog"],
        ["SWA-163", "Event Brackets on Schedule Sub-Page For Public Events Listing and OSS/CMS", "\u2014", "Backlog"],
        ["SWA-137", "Default Event img not in Event Details", "\u2014", "Backlog"],
        ["SWA-136", "[UI Bugs] Card height", "\u2014", "Backlog"],
        ["SWA-128", "[UI Bugs] Use Geist font here [P1]", "\u2014", "Backlog"],
        ["SWA-125", "[UI Bugs] Make the selected font white", "\u2014", "Backlog"],
        ["SWA-121", "[UI Bugs] Inconsistent Card spacing", "\u2014", "Backlog"],
    ])
    doc.add_paragraph()

    add_h1(doc, "P1 \u2014 Nice-to-haves (Medium priority)")
    add_body(doc, "14 tickets. Include as capacity allows.")
    add_table(doc, ["Key", "Summary", "Assignee", "Status"], [
        ["SWA-167", "Add \"Fastbreak\" Badge to Event Cards for Fastbreak Events", "\u2014", "Backlog"],
        ["SWA-166", "One-Time Import of External Events from Jarmal's List into Event Listings", "\u2014", "Backlog"],
        ["SWA-165", "One-Time Import of Compete Legacy Events into Event Listings", "\u2014", "Backlog"],
        ["SWA-164", "BallerTV Live Streaming \u2014 Banner & Live Game Links", "\u2014", "Backlog"],
        ["SWA-146", "Map Size", "\u2014", "Backlog"],
        ["SWA-135", "[UI Bugs] Layout?", "\u2014", "Backlog"],
        ["SWA-133", "[UI Bugs] Make this a carousel", "\u2014", "Backlog"],
        ["SWA-132", "[UI Bugs] Top align", "\u2014", "Backlog"],
        ["SWA-130", "[UI Bugs] MISSING schedule EMPTY STATE", "\u2014", "Backlog"],
        ["SWA-129", "[UI Bugs] MISSING DIVISION EMPTY STATE", "\u2014", "Backlog"],
        ["SWA-126", "[UI Bugs] Add a \"x\" in active state", "\u2014", "Backlog"],
        ["SWA-123", "[UI Bugs] UPDATE TO \"RESULTS PER PAGE\"", "\u2014", "Backlog"],
        ["SWA-122", "[UI Bugs] Search & date icons too small", "\u2014", "Backlog"],
        ["SWA-100", "Fastbreak Ticket / Travel CTAs", "\u2014", "Backlog"],
    ])
    doc.add_paragraph()

    add_h1(doc, "P2 \u2014 Low priority")
    add_body(doc, "10 tickets in scope; not listed. Defer unless P0/P1 are complete and capacity remains.")
    doc.add_paragraph()

    add_h1(doc, "Priority mapping (reference)")
    add_table(doc, ["Jira Priority", "Sprint tier"], [
        ["Highest", "P0 (Must-have)"],
        ["Medium", "P1 (Nice-to-have)"],
        ["Low / other", "P2 (Low priority)"],
    ])

    os.makedirs(os.path.dirname(OUT_PATH), exist_ok=True)
    doc.save(OUT_PATH)
    print(f"Saved to {OUT_PATH}")


if __name__ == "__main__":
    main()
