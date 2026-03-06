from docx import Document
from docx.shared import Pt, RGBColor, Inches, Twips
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

ref_path = "/Users/garretmann/Downloads/Swoop x Club App PRD.docx"
out_path = os.path.join(os.path.dirname(__file__), "..", "prds", "custom-domains-crawl.docx")

doc = Document(ref_path)

for table in list(doc.tables):
    tbl = table._element
    tbl.getparent().remove(tbl)
for p in list(doc.paragraphs):
    p._element.getparent().remove(p._element)

BLUE = RGBColor(0x3D, 0x85, 0xC6)
BEBAS = "Bebas Neue"

LOGO_PATH = os.path.join(os.path.dirname(__file__), "logo.jpg")

def add_title_block(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_picture(LOGO_PATH, width=Inches(2.58), height=Inches(0.35))

    p = doc.add_paragraph()
    run = p.add_run("Custom Domains\nBYOD (Bring Your Own Domain)")
    run.font.name = BEBAS
    run.font.size = Pt(48)
    run.font.bold = True
    run.font.color.rgb = BLUE
    p.style = doc.styles["Title"]

    p = doc.add_paragraph()
    run = p.add_run("Product Requirements")
    run.font.name = BEBAS
    run.font.size = Pt(30)

    p = doc.add_paragraph()
    run = p.add_run("March 2026 DRAFT")
    run.font.size = Pt(10)

    p = doc.add_paragraph()
    run = p.add_run("Confidential")
    run.font.size = Pt(10)
    run.font.bold = True

    for _ in range(3):
        doc.add_paragraph()

def add_h1(doc, text):
    p = doc.add_paragraph(text, style="Heading 1")
    for run in p.runs:
        run.font.name = BEBAS
        run.font.color.rgb = BLUE

def add_h2(doc, text):
    p = doc.add_paragraph(text, style="Heading 2")
    for run in p.runs:
        run.font.name = BEBAS
        run.font.size = Pt(15)

def add_h3(doc, text):
    p = doc.add_paragraph(text, style="Heading 3")

def add_body(doc, text):
    doc.add_paragraph(text)

def add_bold_label(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.bold = True

def add_bullet(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Inches(0.5)
    pf.first_line_indent = Inches(-0.25)
    p.add_run("\u2022  " + text)

def add_numbered_item(doc, num, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Inches(0.5)
    pf.first_line_indent = Inches(-0.25)
    p.add_run(f"{num}.  {text}")

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.bold = True
        run.font.size = Pt(10)

    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(10)

    tbl = table._element
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl.makeelement(qn("w:tblPr"), {})
    borders = tblPr.makeelement(qn("w:tblBorders"), {})
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.makeelement(qn(f"w:{edge}"), {
            qn("w:val"): "single",
            qn("w:sz"): "4",
            qn("w:space"): "0",
            qn("w:color"): "999999",
        })
        borders.append(el)
    tblPr.append(borders)

    return table

# --- Build the document ---

add_title_block(doc)

add_h1(doc, "1. Ship by")
add_body(doc, "ASAP")

add_h1(doc, "2. What & why")
add_body(doc, "Let customers connect their own domain to their site by adding DNS records at their registrar, with copy-paste values and registrar-specific guides so they never need to understand DNS.")

add_h1(doc, "3. Problem")
add_body(doc, "Connecting a custom domain today is manual, error-prone, and drives support tickets. Non-technical customers shouldn\u2019t need to understand DNS to get a working website and email.")

add_h1(doc, "4. What we\u2019re building")
add_body(doc, "Bring-your-own-domain (BYOD) with self-service DNS. The customer enters their domain, we register it on the Vercel project and pull the required DNS config (A record IP + CNAME) from Vercel, then show the customer exactly which records to add (with registrar-specific guides and one-click copy). We store the domain and route traffic via middleware. No nameserver changes, no OAuth flows.")

add_bold_label(doc, "Preconditions:")
add_bullet(doc, "Operator has completed the Operator Standard Sites onboarding flow.")
add_bullet(doc, "Operator is enabled with Custom Sites and has their own domain to use.")

add_h1(doc, "5. Design / UX")
add_bullet(doc, "Form to enter domain.")
add_bullet(doc, "DNS instructions panel with one-click copy for each record value (A record + CNAME).")
add_bullet(doc, "\u201cWhich registrar are you using?\u201d selector that shows registrar-specific step-by-step guides (e.g. GoDaddy, Namecheap).")
add_bullet(doc, "DNS status indicator showing whether records are configured correctly (system checks automatically; customer can return later to see current status).")

add_h1(doc, "6. User flows")
add_h2(doc, "6.1 BYOD: self-service DNS")

add_table(doc,
    ["Step", "Customer does", "System does"],
    [
        ["1", "Opens custom domain / settings in CMS.", "Shows form to enter domain."],
        ["2", "Enters their domain (e.g. mysite.com).", "Validates format; saves to sites.custom_url. Calls Vercel to register domain on project. Calls Vercel to GET config (IP for A record + CNAME record)."],
        ["3", "Selects \u201cWhich registrar are you using?\u201d (e.g. GoDaddy).", "Shows registrar-specific guide + DNS records with copy buttons (CNAME, A)."],
        ["4", "Logs into their registrar, follows directions we provide.", "\u2014"],
        ["5", "Waits for DNS to propagate (minutes to hours).", "Middleware already matches Host to sites.custom_url; site serves when DNS resolves."],
        ["6", "Returns later to check status.", "Runs DNS check and shows status. Domain connected and pointing to the custom site \u2014 tell user this."],
    ]
)

add_body(doc, "")
add_bold_label(doc, "Outcome: Domain connected and pointing to the custom site; customer owns and manages DNS at their registrar.")

add_h1(doc, "7. Architecture overview")

add_body(doc, "Flow (sequential):")
steps = [
    "Operator enabled w/ Custom Sites",
    "Customer opens custom domain settings",
    "Customer enters their domain",
    "Validate format; store in sites.custom_url",
    "Call Vercel: register domain on project",
    "Call Vercel: GET config \u2192 IP for A record + CNAME",
    "Show registrar-specific guide + DNS records with copy (CNAME, A)",
    "Customer follows directions at registrar",
    "DNS propagates",
    "Middleware matches Host \u2192 sites.custom_url \u2192 serve site",
    "Run DNS check and show status",
    "Domain connected and pointing to custom site \u2014 tell user",
]
for i, step in enumerate(steps, 1):
    add_numbered_item(doc, i, step)

add_body(doc, "")
add_bold_label(doc, "End state: sites.custom_url contains the customer\u2019s domain; A record and CNAME are configured by the customer at their registrar; middleware routes incoming requests to the correct site; Vercel handles SSL.")

add_h1(doc, "8. Requirements")

add_h2(doc, "8.1 Domain entry and storage")
add_bold_label(doc, "P0")
add_bullet(doc, "UI for customer to enter their domain; persist in sites.custom_url.")
add_bullet(doc, "Validate domain format before saving.")
add_bullet(doc, "Call Vercel API to register the domain on the Vercel project.")
add_bullet(doc, "Call Vercel API to GET config (returns IP for A record + CNAME record).")

add_h2(doc, "8.2 DNS instructions")
add_bold_label(doc, "P0")
add_bullet(doc, "Display DNS records to add (A record + CNAME) with one-click copy; values come from Vercel config response.")
add_bullet(doc, "Registrar-specific step-by-step guides (e.g. GoDaddy, Namecheap) so customer knows where to add records.")
add_bullet(doc, "DNS verification check runs automatically; show status (pending / active) when customer views the domain settings page. When domain is connected, tell the user.")

add_h2(doc, "8.3 Routing")
add_bold_label(doc, "P0")
add_bullet(doc, "Extend middleware to match incoming request\u2019s base domain against sites.custom_url and serve the correct site data.")

add_h2(doc, "8.4 Cross-cutting")
add_bold_label(doc, "P0")
add_bullet(doc, "We never own the domain; customer always owns it at their registrar.")

add_h1(doc, "9. How we\u2019ll know it worked")
add_bullet(doc, "Customers stop filing support tickets about domain setup.")
add_bullet(doc, "Domains work on first try using the guides we provide.")
add_bullet(doc, "No manual DNS steps required from our team when a customer wants a custom domain.")

add_h1(doc, "10. Out of scope")
add_bullet(doc, "Vercel OAuth flow or customer-facing \u201cConnect Vercel\u201d integration.")
add_bullet(doc, "OpenSRS managed DNS / nameserver-based setup.")
add_bullet(doc, "Domain purchase UI in our app (customer buys at their registrar or Vercel).")
add_bullet(doc, "Domain transfers, renewal management.")
add_bullet(doc, "Billing for domains (customer pays their registrar).")
add_bullet(doc, "Email DNS configuration (MX, SPF, DKIM, DMARC records).")

add_h1(doc, "11. Open questions")
add_bullet(doc, "Which registrars to support in guides first?")

add_h1(doc, "12. Tech / constraints")
add_bullet(doc, "Existing sites table with custom_url column; existing middleware to extend for domain routing.")
add_bullet(doc, "Vercel API: register domain on project + GET config to retrieve A record IP and CNAME value. Vercel auto-provisions SSL for domains added to projects.")
add_bullet(doc, "Preconditions: Operator must have completed Operator Standard Sites onboarding, be enabled with Custom Sites, and have their own domain.")

add_h1(doc, "13. Timeline")
add_table(doc,
    ["Scope", "Est. timeline"],
    [
        ["BYOD + manual DNS guides + middleware", "2\u20133 days"],
    ]
)

doc.save(out_path)
print(f"Saved to {out_path}")
