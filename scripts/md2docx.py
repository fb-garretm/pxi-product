#!/usr/bin/env python3
"""
Convert a Markdown file to a python-docx Document (in memory).
Used by md2docx.py (save to .docx) and md2gdoc.py (upload to Google Docs).
"""
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt


def parse_inline(s):
    """Split text into (text, bold) segments for **bold**."""
    parts = []
    while s:
        m = re.match(r"\*\*([^*]+)\*\*", s)
        if m:
            parts.append((m.group(1), True))
            s = s[m.end() :]
        else:
            idx = s.find("**")
            if idx == -1:
                parts.append((s, False))
                break
            parts.append((s[:idx], False))
            s = s[idx:]
    return parts


def add_paragraph_with_format(doc, line, style=None):
    p = doc.add_paragraph(style=style)
    for text, bold in parse_inline(line.strip()):
        r = p.add_run(text + " " if bold else text)
        r.bold = bold
    return p


def build_document_from_lines(lines):
    """Build a python-docx Document from a list of markdown lines."""
    doc = Document()
    i = 0
    in_code = False
    code_lines = []
    in_table = False
    table_rows = []

    while i < len(lines):
        line = lines[i]
        raw = line.rstrip()

        # Code block
        if raw.startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(raw)
            i += 1
            continue

        # Table: line with |---|
        if re.match(r"^\s*\|[\s\-:]+\|", raw):
            if not in_table and table_rows:
                prev = table_rows
                table_rows = [prev[0]]
                tbl = doc.add_table(rows=1, cols=len(table_rows[0]))
                tbl.style = "Table Grid"
                for ci, cell in enumerate(table_rows[0]):
                    tbl.rows[0].cells[ci].text = cell.strip()
                for row in table_rows[1:] if len(prev) > 1 else []:
                    table_rows.append(row)
                for ri in range(1, len(prev)):
                    row_cells = prev[ri]
                    r = tbl.add_row()
                    for ci, cell in enumerate(row_cells):
                        if ci < len(r.cells):
                            r.cells[ci].text = cell.strip()
                in_table = True
                table_rows = []
            i += 1
            continue

        # Table row
        if raw.strip().startswith("|") and raw.strip().endswith("|"):
            cells = [c.strip() for c in raw.strip().split("|")[1:-1]]
            if not in_table:
                table_rows.append(cells)
            else:
                tbl = doc.tables[-1]
                r = tbl.add_row()
                for ci, cell in enumerate(cells):
                    if ci < len(r.cells):
                        r.cells[ci].text = cell
            i += 1
            continue
        else:
            if table_rows and not in_table and len(table_rows) >= 2:
                sep = table_rows[1]
                if all(re.match(r"^[\s\-]+$", c) for c in sep):
                    header = table_rows[0]
                    tbl = doc.add_table(rows=1, cols=len(header))
                    tbl.style = "Table Grid"
                    for ci, cell in enumerate(header):
                        tbl.rows[0].cells[ci].text = cell.strip()
                    for row in table_rows[2:]:
                        r = tbl.add_row()
                        for ci, cell in enumerate(row):
                            if ci < len(r.cells):
                                r.cells[ci].text = cell.strip()
                else:
                    for row in table_rows:
                        add_paragraph_with_format(doc, " | ".join(row))
                table_rows = []
            in_table = False

        # Headings
        if raw.startswith("# "):
            doc.add_heading(raw[2:].strip(), level=0)
            i += 1
            continue
        if raw.startswith("## "):
            doc.add_heading(raw[3:].strip(), level=1)
            i += 1
            continue
        if raw.startswith("### "):
            doc.add_heading(raw[4:].strip(), level=2)
            i += 1
            continue
        if raw.startswith("#### "):
            doc.add_heading(raw[5:].strip(), level=3)
            i += 1
            continue

        if raw.strip() in ("---", "***", "___"):
            doc.add_paragraph("—" * 20)
            i += 1
            continue

        if re.match(r"^[\-\*]\s+", raw) or re.match(r"^\d+\.\s+", raw):
            add_paragraph_with_format(doc, re.sub(r"^[\-\*]\s+", "", raw), style="List Bullet")
            i += 1
            continue

        if not raw.strip():
            i += 1
            continue

        if re.match(r"^\-?\s*\[[ xX]\]", raw):
            add_paragraph_with_format(doc, raw.strip(), style="List Bullet")
            i += 1
            continue

        add_paragraph_with_format(doc, raw)
        i += 1

    if table_rows and len(table_rows) >= 1:
        header = table_rows[0]
        if len(table_rows) >= 2 and all(re.match(r"^[\s\-]+$", c) for c in table_rows[1]):
            tbl = doc.add_table(rows=1, cols=len(header))
            tbl.style = "Table Grid"
            for ci, cell in enumerate(header):
                tbl.rows[0].cells[ci].text = cell.strip()
            for row in table_rows[2:]:
                r = tbl.add_row()
                for ci, cell in enumerate(row):
                    if ci < len(r.cells):
                        r.cells[ci].text = cell.strip()
        else:
            for row in table_rows:
                add_paragraph_with_format(doc, " | ".join(row))

    return doc


def md_to_docx(md_path):
    """Read markdown from md_path, build doc, save as .docx beside the .md file."""
    md_path = Path(md_path)
    if md_path.suffix.lower() != ".md":
        raise ValueError("Expected a .md file")
    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()
    doc = build_document_from_lines(lines)
    out_path = md_path.with_suffix(".docx")
    doc.save(out_path)
    return out_path


if __name__ == "__main__":
    import sys

    if len(sys.argv) != 2:
        print("Usage: python scripts/md2docx.py <file.md>")
        sys.exit(1)
    out_path = md_to_docx(sys.argv[1])
    print(f"Created: {out_path}")
